"""
Tải dataset bài luận học viên ESL → convert sang .docx → chạy ETL vào DB.

Quy trình:
1. Tải dataset (ELLIPSE từ Kaggle hoặc fallback URL trực tiếp)
2. Mỗi bài luận → tạo 1 file .docx riêng (kèm tiêu đề + tên author)
3. Gọi etl_v3.run_etl() cho từng file → đổ vào 5 bảng

Cách dùng:
    pip install python-docx kaggle requests pandas
    python download_essays_and_etl.py --source ellipse --limit 50
    python download_essays_and_etl.py --source sample        # dùng 5 bài mẫu offline
    python download_essays_and_etl.py --source local --dir ./my_essays
"""

import argparse
import io
import sys
import zipfile
import logging
import urllib.request
from pathlib import Path
import pandas as pd
from docx import Document

sys.path.insert(0, str(Path(__file__).parent))
from etl_v3 import run_etl  # tận dụng pipeline đã có

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
log = logging.getLogger(__name__)

ESSAY_DIR = Path('./student_essays'); ESSAY_DIR.mkdir(exist_ok=True)


# ==========================================================================
# DATASET LOADERS
# ==========================================================================
def load_ellipse(limit: int = 50) -> pd.DataFrame:
    """ELLIPSE — English Language Learner essays (Kaggle).

    Cần đăng ký Kaggle API trước. Hoặc fallback: tải bản mirror trực tiếp.
    """
    try:
        import kaggle
        log.info('Tải ELLIPSE từ Kaggle...')
        kaggle.api.competition_download_files(
            'feedback-prize-english-language-learning',
            path=str(ESSAY_DIR), quiet=False
        )
        zip_path = ESSAY_DIR / 'feedback-prize-english-language-learning.zip'
        with zipfile.ZipFile(zip_path) as z:
            z.extractall(ESSAY_DIR)
        df = pd.read_csv(ESSAY_DIR / 'train.csv')
        df = df.rename(columns={'text_id': 'id', 'full_text': 'text'})
        df['title'] = 'ELL Essay ' + df['id'].astype(str)
        df['author'] = 'Anonymous Student ' + (df.index % 100).astype(str)
        return df[['id', 'title', 'author', 'text']].head(limit)
    except Exception as e:
        log.warning(f'Kaggle fail: {e}. Dùng sample data.')
        return load_sample()


def load_wi_locness(limit: int = 50) -> pd.DataFrame:
    """W&I+LOCNESS từ BEA-2019 (download trực tiếp, không cần auth)."""
    url = 'https://www.cl.cam.ac.uk/research/nl/bea2019st/data/wi+locness_v2.1.bea19.tar.gz'
    log.info(f'Tải W&I+LOCNESS từ {url}...')
    import tarfile, tempfile
    try:
        with tempfile.NamedTemporaryFile(suffix='.tar.gz', delete=False) as tmp:
            urllib.request.urlretrieve(url, tmp.name)
            with tarfile.open(tmp.name) as tar:
                tar.extractall(ESSAY_DIR / 'wi_locness')
    except Exception as e:
        log.warning(f'WI fail: {e}. Dùng sample.'); return load_sample()
    # Parse các file .m2 thành text essays
    rows = []
    for m2_file in (ESSAY_DIR / 'wi_locness').rglob('*.m2'):
        with open(m2_file, encoding='utf-8') as f:
            current_essay = []
            essay_id = 0
            for line in f:
                if line.startswith('S '):
                    current_essay.append(line[2:].strip())
                elif not line.strip() and current_essay:
                    rows.append({
                        'id': f'{m2_file.stem}_{essay_id}',
                        'title': f'W&I Essay {m2_file.stem} #{essay_id}',
                        'author': f'WI Student {essay_id % 50}',
                        'text': '\n'.join(current_essay)
                    })
                    current_essay = []; essay_id += 1
                    if len(rows) >= limit: break
            if len(rows) >= limit: break
        if len(rows) >= limit: break
    return pd.DataFrame(rows)


def load_sample() -> pd.DataFrame:
    """5 bài luận mẫu có lỗi chính tả/ngữ pháp — không cần internet."""
    samples = [
        {
            'id': 'sample_001', 'author': 'Nguyen Van A',
            'title': 'My Hobbies',
            'text': """My hobbies is reading books and playing footbal. I have been collected
stamps since I was eight years old. My favourite book is Harry Potter that
written by JK Rowling. Last summer, I goed to the libary every day to borrow
new books. I think reading make me more inteligent and creativ. In the futur
I want to become a writer like my favorite author. I will write stories that
make people happy and inspired."""
        },
        {
            'id': 'sample_002', 'author': 'Tran Thi B',
            'title': 'The Impact of Technology on Education',
            'text': """Technology have changed education in many ways. Students nowdays can
acces information easily through the internet. However, there are some
disadvantages also. Many student depend on technology too much and they
can not think by themself. Teachers also faces difficulty when their
student use phone in class. In my oppinion, technology is good but we must
use it wisely. We should balance between traditional learning and modern
methods to acheive the best results."""
        },
        {
            'id': 'sample_003', 'author': 'Le Hoang C',
            'title': 'My Summer Vacation',
            'text': """Last summer was the most wonderfull time in my life. My family
and me went to Da Nang beach for one week. The wether was very nice
and the sea was beatiful. We swimmed every morning and ate sea food
in the evening. I also visited Hoi An ancient town which is very
famous tourist atraction. There are many old houses and lanterns
that make the town looks magical at night. I took alot of photos
and bought some souvenirs for my freinds."""
        },
        {
            'id': 'sample_004', 'author': 'Pham Thi D',
            'title': 'Why I Want to Study Abroad',
            'text': """Studying aboard has been my dream since I was a child. I beleive
that studying in foreign country will help me improve my english
and learn about diffrent cultures. My parents was supportive but
they worry about the cost. I have applyed for several scholarship
programs and I hope I can get one. If I succeed, I will study
computer science at a top university. After graduation, I plan to
come back to Vietnam and contribute my knowledge to develop our
country."""
        },
        {
            'id': 'sample_005', 'author': 'Vo Minh E',
            'title': 'Climate Change',
            'text': """Climate change is one of the most serious problem in the world today.
The temperture is rising every year because of green house gas emissions.
Many countries has experience extreme wether such as floods, droughts and
heat waves. Animals are loosing their habitats and some species become
extinct. We need to take action immediatly to save our planet. Each person
can contribute by reducing plastic, planting trees, and using public
transport. Goverments must also implement strict polices to control
pollution and promote renewable energy."""
        },
    ]
    return pd.DataFrame(samples)


def load_local(folder: Path) -> pd.DataFrame:
    """Đọc các file .txt từ folder local của bạn."""
    rows = []
    for i, f in enumerate(folder.glob('*.txt')):
        rows.append({
            'id': f.stem,
            'title': f.stem.replace('_', ' ').title(),
            'author': f'Local Student {i+1}',
            'text': f.read_text(encoding='utf-8', errors='ignore')
        })
    return pd.DataFrame(rows)


# ==========================================================================
# CONVERT TO .DOCX
# ==========================================================================
def text_to_docx(text: str, title: str, out_path: Path):
    """Tạo file .docx có heading + paragraphs."""
    doc = Document()
    doc.add_heading(title, level=1)
    for para in text.split('\n\n'):
        para = para.strip()
        if para:
            doc.add_paragraph(para)
    doc.save(out_path)


# ==========================================================================
# MAIN
# ==========================================================================
def main():
    p = argparse.ArgumentParser()
    p.add_argument('--source', default='sample',
                   choices=['ellipse', 'wi_locness', 'sample', 'local'])
    p.add_argument('--limit', type=int, default=20,
                   help='Số bài luận tối đa (mặc định 20)')
    p.add_argument('--dir', type=str, default='./my_essays',
                   help='Folder chứa .txt khi --source=local')
    p.add_argument('--no-etl', action='store_true',
                   help='Chỉ tạo .docx, không chạy ETL vào DB')
    args = p.parse_args()

    # 1. Load dataset
    if args.source == 'ellipse':
        df = load_ellipse(args.limit)
    elif args.source == 'wi_locness':
        df = load_wi_locness(args.limit)
    elif args.source == 'local':
        df = load_local(Path(args.dir))
    else:
        df = load_sample()
    log.info(f'Đã load {len(df)} bài luận từ {args.source}')

    # 2. Convert sang .docx
    docx_paths = []
    for _, row in df.iterrows():
        path = ESSAY_DIR / f'{row["id"]}.docx'
        text_to_docx(row['text'], row['title'], path)
        docx_paths.append((path, row['author'], row['title']))
    log.info(f'Đã tạo {len(docx_paths)} file .docx tại {ESSAY_DIR}/')

    if args.no_etl:
        log.info('Bỏ qua ETL (--no-etl). Xong.')
        return

    # 3. Chạy ETL từng file vào DB
    success, fail = 0, 0
    for path, author, title in docx_paths:
        try:
            run_etl(str(path), author, title)
            success += 1
        except Exception as e:
            log.error(f'ETL fail {path.name}: {e}')
            fail += 1
    log.info(f'\nKết quả: {success} thành công, {fail} thất bại / tổng {len(docx_paths)}')


if __name__ == '__main__':
    main()
